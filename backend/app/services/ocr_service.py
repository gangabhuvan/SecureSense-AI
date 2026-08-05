"""
ocr_service.py

Unified text extraction service for SecureSense AI.

Supported inputs:
- PDF with embedded text
- Scanned PDF
- PNG / JPG / JPEG / BMP / TIFF images
- DOCX documents
- TXT files

OCR is performed using EasyOCR when required.

Design principles:
- General-purpose and communication-agnostic OCR
- Multi-scale recognition for text of different sizes
- Primary + recovery OCR strategy
- Near-duplicate suppression across OCR passes
- No URL/domain/email/phishing-specific corrections
- No hard-coded assumptions about document content
- Downstream services remain responsible for semantic
  interpretation and entity extraction
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import List

import cv2
import easyocr
import numpy as np
import pdfplumber

from docx import Document
from pdf2image import convert_from_path


# ==========================================================
# OCR / Content Extraction Service
# ==========================================================

class OCRService:
    """
    Extract textual content from supported documents.

    EasyOCR is initialized once and reused across requests.

    OCR remains content-agnostic. It attempts to recover the
    textual representation of the source without making
    assumptions about whether that text represents URLs,
    financial information, phishing content, contact details,
    or any other semantic category.
    """

    IMAGE_EXTENSIONS = {
        ".png",
        ".jpg",
        ".jpeg",
        ".bmp",
        ".tif",
        ".tiff",
    }

    SUPPORTED_EXTENSIONS = {
        ".pdf",
        ".docx",
        ".txt",
        *IMAGE_EXTENSIONS,
    }

    # Minimum amount of embedded PDF text considered useful
    # enough to avoid an unnecessary OCR fallback.
    PDF_TEXT_THRESHOLD = 50

    # Enlargement factor used by the recovery OCR pass.
    OCR_RECOVERY_SCALE = 2.0

    # Minimum normalized token length before substring-based
    # duplicate detection is allowed.
    DUPLICATE_SUBSTRING_MIN_LENGTH = 6

    def __init__(
        self,
    ) -> None:
        """
        Initialize the shared EasyOCR reader.
        """

        self.reader = easyocr.Reader(
            ["en"],
            gpu=False,
        )

    # ======================================================
    # Main Extraction Entry Point
    # ======================================================

    def extract_text(
        self,
        file_path: str,
    ) -> str:
        """
        Detect the input type and extract textual content.
        """

        path = Path(
            file_path
        )

        if not path.exists():
            raise FileNotFoundError(
                f"File not found: {file_path}"
            )

        if not path.is_file():
            raise ValueError(
                f"Path is not a file: {file_path}"
            )

        extension = (
            path.suffix.lower()
        )

        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type: {extension}"
            )

        # --------------------------------------------------
        # PDF
        # --------------------------------------------------

        if extension == ".pdf":

            text = self._extract_from_pdf(
                str(path)
            )

        # --------------------------------------------------
        # DOCX
        # --------------------------------------------------

        elif extension == ".docx":

            text = self._extract_from_docx(
                str(path)
            )

        # --------------------------------------------------
        # TXT
        # --------------------------------------------------

        elif extension == ".txt":

            text = self._extract_from_txt(
                str(path)
            )

        # --------------------------------------------------
        # Images
        # --------------------------------------------------

        elif extension in self.IMAGE_EXTENSIONS:

            text = self._extract_from_image(
                str(path)
            )

        else:

            # Defensive fallback.
            raise ValueError(
                f"Unsupported file type: {extension}"
            )

        return self._normalize_text(
            text
        )

    # ======================================================
    # Text Normalization
    # ======================================================

    @staticmethod
    def _normalize_text(
        text: str,
    ) -> str:
        """
        Normalize unnecessary whitespace while preserving
        extracted textual content.

        No spelling corrections or semantic transformations
        are performed.
        """

        if not text:
            return ""

        normalized_lines: List[str] = []

        for line in text.splitlines():

            cleaned = re.sub(
                r"\s+",
                " ",
                line,
            ).strip()

            if cleaned:
                normalized_lines.append(
                    cleaned
                )

        return "\n".join(
            normalized_lines
        ).strip()

    # ======================================================
    # PDF Extraction
    # ======================================================

    def _extract_from_pdf(
        self,
        pdf_path: str,
    ) -> str:
        """
        Extract text from a PDF.

        Native embedded PDF text is preferred because it is
        normally more accurate than OCR.

        If insufficient embedded text is available, the PDF
        is treated as scanned/image-based and OCR is used.
        """

        pages_text: List[str] = []

        try:

            with pdfplumber.open(
                pdf_path
            ) as pdf:

                for page in pdf.pages:

                    try:

                        page_text = (
                            page.extract_text()
                            or ""
                        )

                    except Exception:

                        page_text = ""

                    page_text = (
                        page_text.strip()
                    )

                    if page_text:

                        pages_text.append(
                            page_text
                        )

        except Exception as exc:

            raise ValueError(
                "Unable to read PDF document."
            ) from exc

        extracted_text = "\n".join(
            pages_text
        ).strip()

        # --------------------------------------------------
        # Prefer embedded PDF text when sufficient content
        # was successfully recovered.
        # --------------------------------------------------

        if (
            len(extracted_text)
            >= self.PDF_TEXT_THRESHOLD
        ):
            return extracted_text

        # --------------------------------------------------
        # Scanned/image-based PDF fallback
        # --------------------------------------------------

        return self._ocr_pdf(
            pdf_path
        )

    # ======================================================
    # Scanned PDF OCR
    # ======================================================

    def _ocr_pdf(
        self,
        pdf_path: str,
    ) -> str:
        """
        Convert scanned PDF pages to images and perform OCR
        on each page independently.
        """

        try:

            images = convert_from_path(
                pdf_path
            )

        except Exception as exc:

            raise ValueError(
                "Unable to convert PDF pages for OCR."
            ) from exc

        pages_text: List[str] = []

        for image in images:

            image_array = np.array(
                image
            )

            if image_array.size == 0:
                continue

            # ------------------------------------------------
            # PIL images generated by pdf2image normally use
            # RGB ordering.
            # ------------------------------------------------

            if len(image_array.shape) == 3:

                if image_array.shape[2] == 4:

                    gray = cv2.cvtColor(
                        image_array,
                        cv2.COLOR_RGBA2GRAY,
                    )

                else:

                    gray = cv2.cvtColor(
                        image_array,
                        cv2.COLOR_RGB2GRAY,
                    )

            else:

                gray = image_array

            page_text = self._run_ocr(
                gray
            )

            if page_text:

                pages_text.append(
                    page_text
                )

        return "\n".join(
            pages_text
        )

    # ======================================================
    # Image OCR
    # ======================================================

    def _extract_from_image(
        self,
        image_path: str,
    ) -> str:
        """
        Extract text from an uploaded image.

        The image may represent any type of communication,
        including screenshots, scanned documents, notices,
        messages, webpages, receipts, posters or photographs.
        """

        image = cv2.imread(
            image_path,
            cv2.IMREAD_UNCHANGED,
        )

        if image is None:
            raise ValueError(
                f"Unable to read image: {image_path}"
            )

        if image.size == 0:
            raise ValueError(
                f"Image contains no readable data: {image_path}"
            )

        # --------------------------------------------------
        # Convert input into grayscale.
        # --------------------------------------------------

        if len(image.shape) == 2:

            gray = image

        elif len(image.shape) == 3:

            channels = image.shape[2]

            if channels == 4:

                gray = cv2.cvtColor(
                    image,
                    cv2.COLOR_BGRA2GRAY,
                )

            elif channels == 3:

                gray = cv2.cvtColor(
                    image,
                    cv2.COLOR_BGR2GRAY,
                )

            else:

                raise ValueError(
                    "Unsupported image channel configuration."
                )

        else:

            raise ValueError(
                "Unsupported image dimensions."
            )

        return self._run_ocr(
            gray
        )

    # ======================================================
    # OCR Image Preparation
    # ======================================================

    @staticmethod
    def _prepare_ocr_image(
        image: np.ndarray,
    ) -> np.ndarray:
        """
        Prepare an image for OCR without applying aggressive
        transformations that could damage characters.

        The method intentionally avoids document-specific
        thresholding, sharpening or semantic preprocessing.
        """

        if image is None or image.size == 0:
            raise ValueError(
                "Cannot prepare an empty image for OCR."
            )

        prepared = image

        # --------------------------------------------------
        # Ensure grayscale
        # --------------------------------------------------

        if len(prepared.shape) == 3:

            prepared = cv2.cvtColor(
                prepared,
                cv2.COLOR_BGR2GRAY,
            )

        # --------------------------------------------------
        # Ensure uint8 representation
        # --------------------------------------------------

        if prepared.dtype != np.uint8:

            prepared = cv2.normalize(
                prepared,
                None,
                0,
                255,
                cv2.NORM_MINMAX,
            ).astype(
                np.uint8
            )

        return prepared

    # ======================================================
    # OCR Pass
    # ======================================================

    def _execute_ocr_pass(
        self,
        image: np.ndarray,
    ) -> List[str]:
        """
        Execute one EasyOCR recognition pass.

        paragraph=False keeps individual OCR detections
        separate so that downstream merging can avoid
        duplicating whole paragraphs between scales.
        """

        try:

            results = self.reader.readtext(
                image,
                detail=1,
                paragraph=False,
            )

        except Exception:

            return []

        detections = []

        for item in results:

            if len(item) != 3:
                continue

            bbox, text, confidence = item

            text = re.sub(
                r"\s+",
                " ",
                str(text),
            ).strip()

            if not text:
                continue

            # Average top Y coordinate
            y = (
                bbox[0][1] + bbox[1][1]
            ) / 2.0

            # Average X coordinate
            x = (
                bbox[0][0] + bbox[3][0]
            ) / 2.0

            # Approximate text height
            text_height = (
                abs(bbox[3][1] - bbox[0][1]) +
                abs(bbox[2][1] - bbox[1][1])
            ) / 2.0

            detections.append(
                (
                    y,
                    x,
                    text_height,
                    text,
                )
            )

        if not detections:
            return []

        # ----------------------------------------
        # Sort top-to-bottom then left-to-right
        # ----------------------------------------

        detections.sort(
            key=lambda t: (
                t[0],
                t[1],
            )
        )

        lines = []

        current_line = []

        current_y = detections[0][0]

        median_height = np.median(
            [d[2] for d in detections]
        )

        LINE_THRESHOLD = max(
            6,
            median_height * 0.6
        )

        for y, x, height, text in detections:

            if abs(
                y - current_y
            ) <= LINE_THRESHOLD:

                current_line.append(
                    (
                        x,
                        text,
                    )
                )

            else:

                current_line.sort(
                    key=lambda t: t[0]
                )

                lines.append(
                    " ".join(
                        value
                        for _, value
                        in current_line
                    )
                )

                current_line = [
                    (
                        x,
                        text,
                    )
                ]

                current_y = y

        if current_line:

            current_line.sort(
                key=lambda t: t[0]
            )

            lines.append(
                " ".join(
                    value
                    for _, value
                    in current_line
                )
            )

        return lines

    # ======================================================
    # OCR Duplicate Comparison
    # ======================================================

    @staticmethod
    def _comparison_form(
        value: str,
    ) -> str:
        """
        Create a normalized representation used only for
        duplicate comparison.

        The original OCR transcription remains unchanged.

        Case, whitespace and punctuation differences are
        ignored for duplicate detection.
        """

        if not value:
            return ""

        normalized = (
            value.lower()
        )

        normalized = re.sub(
            r"[^a-z0-9]+",
            "",
            normalized,
        )

        return normalized

    # ======================================================
    # Near-Duplicate Detection
    # ======================================================

    def _is_near_duplicate(
        self,
        candidate: str,
        existing_values: List[str],
    ) -> bool:
        """
        Determine whether a recovery-pass detection already
        exists in the primary OCR output.

        Conservative matching is used so that genuinely
        different short words are not accidentally removed.
        """

        candidate_key = self._comparison_form(
            candidate
        )

        if not candidate_key:
            return True

        for existing in existing_values:

            existing_key = self._comparison_form(
                existing
            )

            if not existing_key:
                continue

            # ------------------------------------------------
            # Exact normalized match
            # ------------------------------------------------

            if candidate_key == existing_key:
                return True

            shortest_length = min(
                len(candidate_key),
                len(existing_key),
            )

            # ------------------------------------------------
            # Conservative containment check
            #
            # Examples such as:
            #
            # "404 - Page Not Found"
            # "404- Page Not Found"
            #
            # can be treated as equivalent without applying
            # spelling or semantic correction.
            # ------------------------------------------------

            if (
                shortest_length
                >= self.DUPLICATE_SUBSTRING_MIN_LENGTH
            ):

                if (
                    candidate_key in existing_key
                    or existing_key in candidate_key
                ):
                    return True

        return False

    # ======================================================
    # Exact Duplicate Removal
    # ======================================================

    def _remove_exact_duplicates(
        self,
        values: List[str],
    ) -> List[str]:
        """
        Remove exact normalized duplicates while preserving
        the original detection order.
        """

        result: List[str] = []

        seen = set()

        for value in values:

            key = self._comparison_form(
                value
            )

            if not key:
                continue

            if key in seen:
                continue

            seen.add(
                key
            )

            result.append(
                value
            )

        return result

    # ======================================================
    # Multi-Scale OCR
    # ======================================================

    def _run_ocr(
        self,
        image: np.ndarray,
    ) -> str:
        """
        Execute general-purpose multi-scale OCR.

        Strategy:

        1. Original-resolution OCR provides the primary
           transcription.

        2. A 2x enlarged representation provides a recovery
           pass for smaller text that may not have been
           recognized at the original resolution.

        3. Exact and conservative near-duplicate detections
           are removed.

        The recovery pass supplements the primary pass rather
        than replacing it.

        No entity-specific corrections are performed here.
        """

        if image is None or image.size == 0:
            return ""

        prepared = self._prepare_ocr_image(
            image
        )

        # ==================================================
        # Pass 1 — Primary OCR
        # ==================================================

        primary_results = (
            self._execute_ocr_pass(
                prepared
            )
        )

        merged_results = (
            self._remove_exact_duplicates(
                primary_results
            )
        )

        # ==================================================
        # Pass 2 — Recovery OCR
        # ==================================================

        height, width = (
            prepared.shape[:2]
        )

        if height <= 0 or width <= 0:

            return "\n".join(
                merged_results
            )

        recovery_width = max(
            1,
            int(
                round(
                    width
                    * self.OCR_RECOVERY_SCALE
                )
            ),
        )

        recovery_height = max(
            1,
            int(
                round(
                    height
                    * self.OCR_RECOVERY_SCALE
                )
            ),
        )

        enlarged = cv2.resize(
            prepared,
            (
                recovery_width,
                recovery_height,
            ),
            interpolation=cv2.INTER_CUBIC,
        )

        recovery_results = (
            self._execute_ocr_pass(
                enlarged
            )
        )

        recovery_results = (
            self._remove_exact_duplicates(
                recovery_results
            )
        )

        # --------------------------------------------------
        # Add only genuinely new recovery detections.
        # --------------------------------------------------

        for value in recovery_results:

            if self._is_near_duplicate(
                value,
                merged_results,
            ):
                continue

            merged_results.append(
                value
            )

        return "\n".join(
            merged_results
        ).strip()

    # ======================================================
    # DOCX Extraction
    # ======================================================

    def _extract_from_docx(
        self,
        docx_path: str,
    ) -> str:
        """
        Extract textual content from a DOCX document.

        Includes:
        - Standard paragraphs
        - Table cells
        """

        try:

            document = Document(
                docx_path
            )

        except Exception as exc:

            raise ValueError(
                "Unable to read DOCX document."
            ) from exc

        content: List[str] = []

        # --------------------------------------------------
        # Paragraphs
        # --------------------------------------------------

        for paragraph in document.paragraphs:

            text = (
                paragraph.text.strip()
            )

            if text:

                content.append(
                    text
                )

        # --------------------------------------------------
        # Tables
        # --------------------------------------------------

        for table in document.tables:

            for row in table.rows:

                row_values: List[str] = []

                for cell in row.cells:

                    value = (
                        cell.text.strip()
                    )

                    if value:

                        row_values.append(
                            value
                        )

                if row_values:

                    content.append(
                        " | ".join(
                            row_values
                        )
                    )

        return "\n".join(
            content
        )

    # ======================================================
    # TXT Extraction
    # ======================================================

    @staticmethod
    def _extract_from_txt(
        txt_path: str,
    ) -> str:
        """
        Extract text from a UTF-8 plain-text document.
        """

        try:

            with open(
                txt_path,
                "r",
                encoding="utf-8-sig",
            ) as file:

                return file.read()

        except UnicodeDecodeError as exc:

            raise ValueError(
                "Unable to decode TXT file as UTF-8."
            ) from exc

        except OSError as exc:

            raise ValueError(
                "Unable to read TXT file."
            ) from exc


# ==========================================================
# Singleton
# ==========================================================

ocr_service = OCRService()